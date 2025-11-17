import fitz  # PyMuPDF
from flask import Flask, request, render_template, jsonify, send_file, send_from_directory
from flask.cli import load_dotenv
from flask_cors import CORS  # For cross-origin requests
import os
from docx import Document
import io
from werkzeug.utils import secure_filename
import traceback
import google.generativeai as genai

load_dotenv()

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app, supports_credentials=True)
CORS(app)  # Enable CORS for all routes

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf'}
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024  # 2MB limit

# Initialize Cohere client (replace with your actual API key)
API_KEY = os.getenv("api_key")

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel(model_name="gemini-1.5-flash-latest")


# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_pdf(pdf_path):
    """Extract text from PDF using PyMuPDF"""
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        raise

@app.route('/')
def index():
    return render_template('signin.html')

@app.route('/signin')
def signin():
    return render_template('signin.html')

@app.route('/signup')
def signup():
    return render_template('signup.html')

@app.route('/dashboard')
def dashboard():
    return render_template('index.html')

@app.route('/static/<path:filename>')
def static_files(filename):
    return send_from_directory(app.static_folder, filename)

@app.route('/generate_cover_letter', methods=['POST'])
def generate_cover_letter():
    filepath = None
    try:
        print("Request received!")

        if 'cv' not in request.files:
            print("No file in request")
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files['cv']
        job = request.form.get('job', '').strip()
        company = request.form.get('company', '').strip()
        company_desc = request.form.get('company_desc', '').strip()

        print("File:", file.filename)
        print("Job:", job)
        print("Company:", company)

        # Check if file is allowed
        if not file or file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        if not allowed_file(file.filename):
            return jsonify({"error": "Only PDF files are allowed"}), 400

        # Save file temporarily
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        print(f"File saved to {filepath}")

        # Extract PDF
        cv_text = extract_pdf(filepath)
        print("Extracted CV Text:", cv_text[:200], "...")  # print preview


        prompt = f"""
        Write a professional, tailored cover letter based on the following information:

        Applicant's CV:
        {cv_text}

        Target Position: {job}
        Target Company: {company}
        Company Description: {company_desc}

        The cover letter should:
        1. Begin with a proper greeting (e.g., “Dear Hiring Manager,” or the hiring manager’s name if provided).
        2. Structure the letter into 3–4 concise paragraphs:
            1. Intro: State the role you’re applying for, your background, and enthusiasm.
            2. Body (1–2 paragraphs): Highlight your most relevant skills, experiences, and projects from the resume, showing measurable impact where possible.
            3. Closing: Express motivation to join the company and summarize what you will contribute.
        3. Use details from the candidate’s resume accurately — include correct education, work experience, and skills.
        4. Align the candidate’s strengths with the job description and the company’s values/mission.
        5. Maintain a professional, confident, and concise tone. Avoid repetition, filler phrases, or generic statements.
        6. Do not use bullet points.
        7. Do not include addresses, phone numbers, email, LinkedIn, GitHub, or dates.
        8. End with the candidate’s full name only (no “Sincerely,” or extra notes).
        9. Format the output as a clean, professional cover letter that can be sent directly online.

        Cover Letter:
        """

        response = model.generate_content(prompt)


        cover_letter = response.text.strip()

        print("Cover letter:", cover_letter)

        return jsonify({
            "success": True,
            "cover_letter": cover_letter,
            "job": job,
            "company": company
        }), 200, {'Content-Type': 'application/json'}

    except Exception as e:
        print("Exception occurred:")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    
    finally:
        if filepath and os.path.exists(filepath):
            os.remove(filepath)

@app.route('/download_cover_letter', methods=['POST'])
def download_cover_letter():
    try:
        data = request.get_json()
        cover_letter = data.get('cover_letter', '')
        
        doc = Document()
        for line in cover_letter.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph('')
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name='cover_letter.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print("Download error:", str(e))
        return jsonify({"error": "Failed to generate document"}), 500

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000, debug=True)
